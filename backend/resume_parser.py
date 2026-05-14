"""
Resume Parser Module
Extracts structured data from DOCX and PDF files.
Uses magic-byte detection so files work regardless of extension.
"""
import re
import os
from pathlib import Path
from typing import Dict, Any, Tuple
import uuid
from datetime import datetime

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError as DocxPackageNotFoundError
except ImportError:
    Document = None
    DocxPackageNotFoundError = Exception

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


# ── Constants ─────────────────────────────────────────────────────────────────

US_STATE_ABBREVS = {
    'AL','AK','AZ','AR','CA','CO','CT','DE','FL','GA','HI','ID','IL','IN','IA',
    'KS','KY','LA','ME','MD','MA','MI','MN','MS','MO','MT','NE','NV','NH','NJ',
    'NM','NY','NC','ND','OH','OK','OR','PA','RI','SC','SD','TN','TX','UT','VT',
    'VA','WA','WV','WI','WY','DC'
}

US_STATE_NAMES = [
    'Alabama','Alaska','Arizona','Arkansas','California','Colorado','Connecticut',
    'Delaware','Florida','Georgia','Hawaii','Idaho','Illinois','Indiana','Iowa',
    'Kansas','Kentucky','Louisiana','Maine','Maryland','Massachusetts','Michigan',
    'Minnesota','Mississippi','Missouri','Montana','Nebraska','Nevada',
    'New Hampshire','New Jersey','New Mexico','New York','North Carolina',
    'North Dakota','Ohio','Oklahoma','Oregon','Pennsylvania','Rhode Island',
    'South Carolina','South Dakota','Tennessee','Texas','Utah','Vermont',
    'Virginia','Washington','West Virginia','Wisconsin','Wyoming'
]

INDIAN_CITIES = [
    'Mumbai','Delhi','Bangalore','Bengaluru','Hyderabad','Chennai','Kolkata',
    'Pune','Ahmedabad','Jaipur','Lucknow','Kanpur','Nagpur','Visakhapatnam',
    'Vizag','Indore','Thane','Bhopal','Patna','Vadodara','Ghaziabad','Ludhiana',
    'Coimbatore','Agra','Madurai','Nashik','Vijayawada','Surat','Rajkot',
    'Amritsar','Meerut','Jabalpur','Warangal','Ranchi','Guwahati','Chandigarh',
    'Thiruvananthapuram','Kochi','Cochin','Noida','Gurugram','Gurgaon',
    'Faridabad','Navi Mumbai','Mysore','Mysuru','Bhubaneswar','Raipur',
    'Dehradun','Jodhpur','Tiruchirappalli','Salem','Guntur','Tirupati',
    'Nellore','Kakinada','Rajahmundry','Khammam','Karimnagar','Nizamabad',
    'Secunderabad','Aurangabad','Solapur','Kolhapur','Hubli','Mangalore',
    'Udupi','Belgaum','Dharwad','Shimla','Jammu','Srinagar','Gwalior',
    'Ujjain','Bhilai','Bokaro','Dhanbad','Jamshedpur','Bilaspur'
]

JOB_TITLES = [
    'Software Engineer','Senior Software Engineer','Junior Software Engineer',
    'Staff Software Engineer','Principal Software Engineer',
    'Full Stack Developer','Full Stack Engineer',
    'Frontend Developer','Frontend Engineer',
    'Backend Developer','Backend Engineer',
    'Web Developer','Web Designer',
    'Data Engineer','Senior Data Engineer','Data Scientist','Data Analyst',
    'Business Intelligence Analyst','BI Developer','Analytics Engineer',
    'Machine Learning Engineer','ML Engineer','AI Engineer',
    'Deep Learning Engineer','NLP Engineer','Computer Vision Engineer',
    'DevOps Engineer','Site Reliability Engineer','SRE',
    'Cloud Engineer','Cloud Architect','Platform Engineer',
    'Infrastructure Engineer','Systems Engineer','Systems Administrator',
    'Android Developer','iOS Developer','Mobile Developer',
    'React Native Developer','Flutter Developer',
    'Python Developer','Java Developer','JavaScript Developer',
    'Node.js Developer','React Developer','Angular Developer',
    'Vue Developer','PHP Developer','.NET Developer','C++ Developer',
    'QA Engineer','Quality Assurance Engineer','Test Engineer',
    'Automation Engineer','SDET','Manual Tester',
    'Product Manager','Project Manager','Program Manager',
    'Marketing Manager','Project Marketing Manager',
    'Digital Marketing Manager','Social Media Manager','Brand Manager',
    'Growth Manager','Content Manager','SEO Manager',
    'Scrum Master','Agile Coach','Engineering Manager',
    'Technical Lead','Tech Lead','Team Lead',
    'Solutions Architect','Software Architect','Enterprise Architect',
    'Technical Architect',
    'Business Analyst','System Analyst','Functional Analyst',
    'Database Administrator','DBA','Database Developer',
    'Security Engineer','Cybersecurity Analyst','Penetration Tester',
    'Network Engineer','Network Administrator',
    'UI/UX Designer','UX Designer','UI Designer','Product Designer',
    'Embedded Systems Engineer','Firmware Engineer','VLSI Engineer',
    'Hardware Engineer','Electronics Engineer',
    'Blockchain Developer','Game Developer',
    'Software Developer','Application Developer','Technology Analyst',
    'Associate Engineer','Senior Engineer','Lead Engineer',
    'Graduate Engineer','Trainee Engineer',
    'Consultant','Marketing Consultant','Web Consultant',
    'Social Media Consultant','Strategy Consultant',
    'Digital Marketing Specialist','SEO Specialist','Content Strategist',
]

SKILL_KEYWORDS = [
    'python','java','javascript','typescript','react','angular','vue',
    'node','node.js','express','django','flask','fastapi','spring',
    'c++','c#','ruby','php','golang','go','rust','swift','kotlin',
    'scala','r','matlab','sql','mysql','postgresql','mongodb',
    'redis','elasticsearch','cassandra','oracle','mssql','sqlite',
    'azure','aws','gcp','docker','kubernetes','terraform','ansible',
    'linux','git','gitlab','github','jenkins','jira','confluence',
    'html','css','tailwind','bootstrap','sass','rest api','graphql',
    'microservices','devops','ci/cd','agile','scrum',
    'machine learning','deep learning','tensorflow','pytorch','keras',
    'pandas','numpy','scikit-learn','spark','hadoop','kafka',
    'power bi','tableau','excel','salesforce','sap',
    'android','ios','react native','flutter','xamarin',
    'selenium','junit','pytest','jest','cypress',
    'unity','unreal','blender','figma','sketch',
    'blockchain','solidity','web3','ethereum',
    'photoshop','lightroom','davinci resolve','adobe',
    'google analytics','google ads','facebook ads','seo','sem',
    'hubspot','hootsuite','sprout social',
]


# ── File Type Detection ───────────────────────────────────────────────────────

def detect_file_type(filepath: str) -> str:
    """
    Detect the real file type using magic bytes, not file extension.
    Returns 'pdf', 'docx', 'doc', or 'unknown'.
    """
    with open(filepath, 'rb') as f:
        header = f.read(8)

    # PDF: starts with %PDF
    if header[:4] == b'%PDF':
        return 'pdf'

    # DOCX/XLSX/PPTX (ZIP-based Office): starts with PK\x03\x04
    if header[:4] == b'PK\x03\x04':
        return 'docx'

    # Old .doc (OLE2 Compound): starts with D0 CF 11 E0
    if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        return 'doc'

    return 'unknown'


# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(filepath: str) -> str:
    if not pdfplumber:
        raise ImportError("pdfplumber not installed.")
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                except Exception:
                    continue
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    return "\n".join(text_parts)


def extract_text_from_docx(filepath: str) -> str:
    if not Document:
        raise ImportError("python-docx not installed.")
    doc = Document(filepath)
    text_parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    return "\n".join(text_parts)


def extract_text_from_doc(filepath: str) -> str:
    """Extract text from old binary .doc using textract or antiword fallback."""
    try:
        import subprocess
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, Exception):
        pass

    # Fallback: try reading as docx anyway (some .doc files are actually docx)
    try:
        return extract_text_from_docx(filepath)
    except Exception:
        pass

    # Last resort: raw byte extraction of readable text
    try:
        with open(filepath, 'rb') as f:
            raw = f.read()
        # Extract ASCII strings of length > 4
        strings = re.findall(rb'[ -~]{4,}', raw)
        text = '\n'.join(s.decode('ascii', errors='ignore') for s in strings)
        return text
    except Exception as e:
        raise Exception(f"Could not extract text from .doc file: {str(e)}")


def extract_text(filepath: str) -> str:
    """
    Smart text extraction: detects real file type by magic bytes,
    not just extension. Handles misnamed files gracefully.
    """
    real_type = detect_file_type(filepath)
    ext = Path(filepath).suffix.lower()

    print(f"[PARSER] File: {os.path.basename(filepath)}, Extension: {ext}, Real type: {real_type}", flush=True)

    # Use real type for extraction
    if real_type == 'pdf':
        return extract_text_from_pdf(filepath)
    elif real_type == 'docx':
        return extract_text_from_docx(filepath)
    elif real_type == 'doc':
        return extract_text_from_doc(filepath)
    else:
        # Unknown magic bytes — try by extension as fallback
        if ext == '.pdf':
            return extract_text_from_pdf(filepath)
        elif ext in ['.docx']:
            return extract_text_from_docx(filepath)
        elif ext in ['.doc']:
            return extract_text_from_doc(filepath)
        else:
            raise ValueError(f"Cannot read file: unrecognised format (ext={ext}, magic={real_type})")


# ── Field Extractors ──────────────────────────────────────────────────────────

def extract_email(text: str) -> str:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else ""


def extract_phone(text: str) -> str:
    patterns = [
        r'\+?1?\s*\(?(\d{3})\)?[-.\s](\d{3})[-.\s](\d{4})',
        r'\b\d{10}\b',
        r'\+\d[\d\s\-\(\)]{8,14}\d',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            m = matches[0]
            if isinstance(m, tuple):
                return ''.join(m)
            digits = re.sub(r'\D', '', m)
            if 7 <= len(digits) <= 15:
                return digits
    return ""


def extract_name(text: str, filename: str) -> str:
    # Try filename pattern: digits_digits_digits_Name.ext
    filename_no_ext = os.path.splitext(filename)[0]
    parts = filename_no_ext.split('_')
    if len(parts) >= 4:
        name = ' '.join(parts[3:]).strip()
        if name and re.match(r'^[A-Za-z\s\.\-]+$', name):
            return name

    # Try first lines of text
    skip = re.compile(
        r'^\s*(resume|curriculum|cv|page|address|phone|email|mobile|contact|'
        r'objective|summary|profile|linkedin|github|http|www|date|\d)',
        re.IGNORECASE
    )
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:8]:
        if skip.match(line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(re.match(r'^[A-Za-z\.\-]+$', w) for w in words):
            return line
    return lines[0][:50] if lines else "Unknown"


def extract_name_from_filename(filename: str) -> str:
    filename_no_ext = os.path.splitext(filename)[0]
    parts = filename_no_ext.split('_')
    if len(parts) >= 4:
        name = ' '.join(parts[3:]).replace('_', ' ').strip()
        if re.match(r'^[A-Za-z\s\.\-]+$', name):
            return name
    return ""


def extract_location(text: str) -> Tuple[str, str]:
    """Search only top 40 lines to avoid picking up skill/project text."""
    lines = text.split('\n')
    search_lines = lines[:40]
    search_text = '\n'.join(search_lines)

    city, state = '', ''

    # Pattern 1: "City, ST" US format (most reliable)
    pattern_us = re.compile(r'\b([A-Z][a-zA-Z\s]{1,20}),\s*([A-Z]{2})\b')
    for line in search_lines[:25]:
        for m in pattern_us.finditer(line):
            pot_city = m.group(1).strip()
            pot_state = m.group(2).strip()
            if (pot_state in US_STATE_ABBREVS and
                    len(pot_city.split()) <= 3 and
                    not any(c.isdigit() for c in pot_city) and
                    len(pot_city) > 2):
                city = pot_city
                state = pot_state
                return city, state

    # Pattern 2: Zip code + City, ST (e.g. "33186 Miami, FL")
    zip_city = re.compile(r'\d{5}\s+([A-Z][a-zA-Z\s]{1,20}),\s*([A-Z]{2})\b')
    for line in search_lines[:15]:
        m = zip_city.search(line)
        if m:
            pot_city = m.group(1).strip()
            pot_state = m.group(2).strip()
            if pot_state in US_STATE_ABBREVS:
                return pot_city, pot_state

    # Pattern 3: Full US state name
    for state_name in US_STATE_NAMES:
        pattern = re.compile(
            rf'\b([A-Z][a-zA-Z\s]{{1,20}}),\s*{re.escape(state_name)}\b',
            re.IGNORECASE
        )
        m = pattern.search(search_text)
        if m:
            city = m.group(1).strip()
            state = state_name
            return city, state

    # Pattern 4: Indian cities
    for city_name in INDIAN_CITIES:
        pattern = re.compile(rf'\b{re.escape(city_name)}\b', re.IGNORECASE)
        m = pattern.search(search_text)
        if m:
            city = city_name.title()
            idx = search_text.lower().find(city_name.lower())
            nearby = search_text[max(0, idx - 30):idx + 80]
            state_m = re.search(r',\s*([A-Z][a-zA-Z\s]{2,20})\b', nearby)
            if state_m:
                state = state_m.group(1).strip()
            return city, state

    return city, state


def _clean_title(raw: str) -> str:
    """Remove trailing location/contract/date noise from a title string."""
    # Remove "(Contract)", "(Full-time)" etc.
    raw = re.sub(r'\s*\((?:contract|full.?time|part.?time|temp|intern|remote)\)\s*', '', raw, flags=re.IGNORECASE)
    # Remove trailing location "Seattle, WA" or "Seattle,WA" — single word city + 2-letter state
    raw = re.sub(r'[\s]*[A-Z][a-z]+,\s*[A-Z]{2}\s*$', '', raw)
    # Remove trailing dates
    raw = re.sub(r'\s+\d{4}.*$', '', raw)
    return raw.strip()


def extract_job_title(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # Step 1: Exact match in first 15 lines
    for line in lines[:15]:
        for title in JOB_TITLES:
            if line.lower() == title.lower():
                return title

    # Step 2: Line starts with a known title (with possible prefix)
    for line in lines[:15]:
        for title in JOB_TITLES:
            if re.match(
                rf'^(senior|sr\.?|junior|jr\.?|lead|principal|staff|associate|'
                rf'mid|mid-level|entry\s+level)?\s*{re.escape(title)}',
                line, re.IGNORECASE
            ):
                return _clean_title(line)[:80]

    # Step 3: Objective/Title label
    label_pattern = re.compile(
        r'(?:objective|career\s+objective|job\s+title|title|position|role|'
        r'designation|current\s+role)\s*[:\-]\s*(.+)',
        re.IGNORECASE
    )
    for line in lines[:30]:
        m = label_pattern.search(line)
        if m:
            extracted = _clean_title(m.group(1))
            if 3 < len(extracted) < 80:
                return extracted

    # Step 4: Scan first 800 chars for known title
    first_chunk = text[:800]
    for title in JOB_TITLES:
        if re.search(rf'\b{re.escape(title)}\b', first_chunk, re.IGNORECASE):
            return title

    # Step 5: Full text scan
    for title in JOB_TITLES:
        if re.search(rf'\b{re.escape(title)}\b', text, re.IGNORECASE):
            return title

    return 'Not Specified'


def extract_work_authorization(text: str) -> str:
    auth_map = {
        'US Citizen': ['us citizen', 'u.s. citizen', 'american citizen'],
        'Green Card': ['green card', 'permanent resident', 'lawful permanent'],
        'H-1B': ['h-1b', 'h1b', 'h1-b'],
        'L-1': ['l-1 visa', 'l1 visa'],
        'O-1': ['o-1 visa', 'o1 visa'],
        'F-1 OPT': ['f-1', 'f1 visa', 'optional practical training', ' opt '],
        'TN Visa': ['tn visa', 'trade nafta'],
        'EAD': ['employment authorization', ' ead '],
        'Open to Sponsorship': ['willing to sponsor', 'require sponsorship'],
    }
    text_lower = text.lower()
    for auth_type, keywords in auth_map.items():
        for kw in keywords:
            if kw in text_lower:
                return auth_type
    return 'Not Specified'


def extract_skills(text: str) -> str:
    found_skills = set()
    text_lower = text.lower()

    # Scan full text for skill keywords
    for skill in SKILL_KEYWORDS:
        if re.search(rf'\b{re.escape(skill)}\b', text_lower):
            found_skills.add(skill.title())

    return ', '.join(sorted(found_skills)) if found_skills else 'Not Specified'


# ── Main Parser ───────────────────────────────────────────────────────────────

def parse_resume(filepath: str) -> Dict[str, Any]:
    try:
        filename = os.path.basename(filepath)

        # Smart extraction — handles misnamed files (PDF saved as .docx etc.)
        text = extract_text(filepath)

        if not text or not text.strip():
            raise ValueError("No text could be extracted from the file")

        city, state = extract_location(text)

        applicant_data = {
            'applicantId': str(uuid.uuid4()),
            'applicantName': extract_name_from_filename(filename) or extract_name(text, filename),
            'emailAddress': extract_email(text),
            'mobileNumber': extract_phone(text),
            'city': city,
            'state': state,
            'applicantStatus': 'Active',
            'jobTitle': extract_job_title(text),
            'ownership': 'Internal',
            'workAuthorization': extract_work_authorization(text),
            'source': 'Resume Upload',
            'createdBy': 'System',
            'createdOn': datetime.now().isoformat(),
            'techSkills': extract_skills(text),
            'resumeText': text,
        }

        return applicant_data

    except Exception as e:
        raise Exception(f"Error parsing resume: {str(e)}")


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        for key, value in result.items():
            if key != 'resumeText':
                print(f"  {key}: {value}")